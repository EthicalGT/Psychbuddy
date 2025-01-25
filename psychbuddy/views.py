import pymysql.cursors
import yagmail
import random
import pymysql
import base64
import socket
import requests
import shutil
import os
import time
import subprocess
from datetime import date, datetime
from django.http import HttpResponse,HttpResponseRedirect,JsonResponse
from docx import Document
from django.shortcuts import render, redirect
from django.http import JsonResponse
from .emotiondetector import detect_emo
from django.contrib import messages
from django.views.decorators.csrf import csrf_exempt

global emotion

# Home page
def index(request):
    return render(request, 'index.html')

# Email verification page
def email(request):
    if (request.method == 'POST'):
        if request.POST.get('signup')=="signup":
            return render(request, 'emailVerify.html')
        else:
            return render(request,'signin.html')
            print("Not yet created")
    return render(request, 'index.html')

#sign up page:
def signUp(req):
    return render(req,"signup.html")


# OTP sending and saving
def OTP(request):
    if request.method == "POST":
        mobile = request.POST.get('mobile')
        email = request.POST.get('email')
        request.session['emailSess']=email
        request.session['mobileSess']=mobile
        
        conn = pymysql.connect(
            host='localhost',
            user='root',
            password='',
            database='psychbuddy'
        )
        try:
            yag = yagmail.SMTP('mypyschbuddy@gmail.com', 'pewcrroshhqeqalf')
            otp = str(random.randint(100000, 999999))
            subject = 'Psychbuddy Email Verification'
            body = f'Dear user, your OTP for Psychbuddy verification is {otp}'
            yag.send(to=email, subject=subject, contents=body)
            with conn.cursor() as cursor:
                sql = "INSERT INTO otp_verification (mobileno, email, otp, date, time) VALUES (%s, %s, %s, %s, %s)"
                cursor.execute(sql, (mobile, email, otp, date.today(), datetime.now().time()))
                conn.commit()
            return render(request, "verifyOTP.html", {'res': "OTP Sent!", 'email': request.session.get('emailSess'), 'mobile': request.session.get('mobileSess')})
        except Exception as e:
            print(f"Error: {e}")
            conn.rollback()
            return render(request, "emailVerify.html", {'res': "Failed to send OTP. Please try again.", 'email': email, 'mobile': mobile})
        finally:
            conn.close()

    return render(request, "emailVerify.html", {'res': "Failed to send OTP. Please try again."})

# OTP verification logic
def otpVerify(request):
    otp = request.POST.get('tb1')
    email=request.session.get('emailSess')
    mobile=request.session.get('mobileSess')
    conn = pymysql.connect(host='localhost', user='root', password='', database='psychbuddy')
    try:
        with conn.cursor() as cursor:
            cursor.execute("SELECT otp FROM otp_verification WHERE otp = %s", (otp,))
            result = cursor.fetchone()

            if str(result and otp) == str(result[0]):
                print(result[0])
                print("Correct OTP")
                return render(request, "signup.html", {'res': 'OTP Verified Go for SignUp.','email': email, 'mobileno': mobile})
            else:
                print("Incorrect OTP")
                return render(request,"verifyOTP",{'res':'Incorrect OTP Please Try Again!'})
    except Exception as e:
        print(f"Error: {e}")
        return render(request, "verifyOTP.html", {'res': 'Error occurred during verification!'})
    finally:
        conn.close()

    return render(request, "verifyOTP.html")

#signup
def signupUser(req):
    if req.method == 'POST':
        fullname = req.POST.get('tb1')
        gender = req.POST.get('rb')
        dob = req.POST.get('tb2')
        dob2 = str(req.POST.get('tb2'))
        year = int(dob2[0:4])
        age = 2024 - year  
        mobile = str(req.session.get('mobileSess'))
        myemail = str(req.session.get('emailSess'))
        addr = req.POST.get('tb3')
        pwd = req.POST.get('tb4')
        status=0
        encodedpwd = base64.b64encode(pwd.encode())
        password = encodedpwd.decode()

        try:
            conn = pymysql.connect(
                host='localhost',
                user='root',
                password='',
                database='psychbuddy'
            )

            with conn.cursor() as cursor:
                check_sql = "SELECT email FROM otp_verification WHERE email = %s"
                cursor.execute(check_sql, (myemail,))
                result = cursor.fetchone()
                
                if result is None:
                    return render(req, "signup.html", {'res': "Email Does not Matches!"})

                sql = "INSERT INTO users (fullname, gender, dob, age, mobileno, email, address, password,login_status) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)"
                cursor.execute(sql, (fullname, gender, dob, age, mobile, myemail, addr, password, status))
                conn.commit()

            print("Data Saved in DB")
            return render(req, "signin.html", {'res': "User Registered Successfully!"})

        except Exception as e:
            print(f"Error: {e}")
            return render(req, "signup.html", {'res': "Error Occurred!"})

        finally:
            conn.close()  

    return render(req, "signup.html")

def signin(req):
    if req.method=='POST':
        email=str(req.POST.get('username'))
        pwd=str(req.POST.get('pass'))
        mypwd = base64.b64encode(pwd.encode()).decode('utf-8')
        status=1
        conn=pymysql.connect(host='localhost',user='root',password='',database='psychbuddy',cursorclass=pymysql.cursors.DictCursor)
        try:
            with conn.cursor() as cursor:
                sql_query = "SELECT * FROM users WHERE email=%s AND password=%s"
                cursor.execute(sql_query, (email, mypwd))
                rows = cursor.fetchall()
                if rows:
                    sql2="UPDATE users set login_status = case when email = %s then 1 else 0 end"
                    cursor.execute(sql2,(email,))
                    conn.commit()
                    return render(req,'homepage.html')
                else:
                    return HttpResponse("<html><body><script>alert('Invalid Credentials Entered. Please Try After Some Time!'); window.location='signin';</script></body></html>")
        except Exception as e:
            print(e)
        finally:
            conn.close()
    return render(req,'signin.html')

def chatForum(request):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password="",
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
    msgs = []
    msgs2=[]

    try:
        with conn.cursor() as cr:
            sql = "select users.fullname, community_forum.date, community_forum.time, community_forum.msg, community_forum.administrator from users join community_forum on community_forum.email = users.email where users.login_status = %s order by community_forum.date asc, community_forum.time asc"
            cr.execute(sql, (1,))
            res = cr.fetchall()
            if res:
                for data in res:
                    admin=None
                    if data['administrator']:
                        admin="Admin"
                    else:
                        admin="User"
                    msgs.append({
                        'msg': data['msg'],
                        'cdate': data['date'],
                        'ctime': data['time'],
                        'name': data['fullname'],
                        'admin': admin
                    })
    except Exception as e:
        print(f"Error: {e}")
    
    try:
        with conn.cursor() as cr:
            sql = "select users.fullname, community_forum.date, community_forum.time, community_forum.msg, community_forum.administrator from users join community_forum on community_forum.email = users.email where users.login_status = %s order by community_forum.date asc, community_forum.time asc"
            cr.execute(sql, (0,))
            res = cr.fetchall()
            if res:
                for data in res:
                    admin=None
                    if data['administrator']:
                        admin="Admin"
                    else:
                        admin="User"
                    msgs2.append({
                        'msg': data['msg'],
                        'cdate': data['date'],
                        'ctime': data['time'],
                        'name': data['fullname'],
                        'admin': admin
                    })
    except Exception as e:
        print(f"Error: {e}")
    finally:
        conn.close()

    if request.method == 'POST':
        msg = request.POST.get('msg')
        email = None
        conn = pymysql.connect(
            host='localhost',
            user='root',
            password='',
            database='psychbuddy',
            cursorclass=pymysql.cursors.DictCursor
        )
        try:
            with conn.cursor() as cr:
                sql = "SELECT email FROM users WHERE login_status = %s"
                cr.execute(sql, (1,))
                rw = cr.fetchone()
                if rw:
                    email = rw['email']
                else:
                    return JsonResponse({'success': False, 'error': 'Email not found'}, status=400)
                sql2 = "INSERT INTO community_forum (msg, email, date, time, administrator) VALUES (%s, %s, %s, %s, %s)"
                cr.execute(sql2, (msg, email, date.today(), datetime.now().time(), None))
                conn.commit()
                return JsonResponse({'success': True, 'msg': msg})
        except Exception as e:
            print(e)
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
        finally:
            conn.close()
    return render(request, "forum.html", {'info': msgs,'info2':msgs2})


def check_connection():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=5)
        return True
    except OSError:
        return False

def news_view(request):
    if check_connection():
        api_key = 'ff76790272f843409058cf099cc785be'
        url = f'https://newsapi.org/v2/everything?q=mental%20health&apiKey={api_key}'
        
        try:
            response = requests.get(url)
            news = []

            if response.status_code == 200:
                data = response.json()
                articles = data.get('articles', [])
                print("Number of articles fetched:", len(articles))

                for article in articles:
                    image_url = article.get('urlToImage')
                    news.append({
                        'title': article['title'],
                        'link': article['url'],
                        'description': article['description'],
                        'image_url': image_url,
                        'published': datetime.strptime(article['publishedAt'], '%Y-%m-%dT%H:%M:%SZ')
                    })
                print("Message -> News Added.")
            else:
                print(f"Failed to fetch news: {response.status_code}")
        
        except requests.RequestException as e:
            print(f"An error occurred while fetching news: {e}")
    
    else:
        print("Not connected")

    return render(request, 'news.html', {'news': news if 'news' in locals() else []})

#therapy session views-------------------------------------------------------------------------->

def individual_therapy(req):
    conn=pymysql.connect(host='localhost',user='root',password='',database='psychbuddy',cursorclass=pymysql.cursors.DictCursor)
    dataset=[]
    therapy=None
    try:
        with conn.cursor() as cr:
            cr.execute("select * from sessions where sess_category=%s",('individual therapy',))
            res=cr.fetchall()
            if res:
                for data in res:
                    therapy=data['sess_category']
                    dataset.append({
                        'name':data['sess_name'],
                        'category':data['sess_category'],
                        'path':data['sess_path'],
                        'desc':data['sess_desc']
                    })
                #print(dataset)
                return render(req,'sessionpage.html',{'myset': dataset,'therapy':therapy})
            else:
                print("no data found")
    except Exception as e:
        print(f"Error->{e}")
    finally: 
        conn.close()


def group_therapy(req):
    conn=pymysql.connect(host='localhost',user='root',password='',database='psychbuddy',cursorclass=pymysql.cursors.DictCursor)
    dataset=[]
    therapy=None
    try:
        with conn.cursor() as cr:
            cr.execute("select * from sessions where sess_category=%s",('Group therapy',))
            res=cr.fetchall()
            if res:
                for data in res:
                    therapy=data['sess_category']
                    dataset.append({
                        'name':data['sess_name'],
                        'category':data['sess_category'],
                        'path':data['sess_path'],
                        'desc':data['sess_desc']
                    })
                #print(dataset)
                return render(req,'sessionpage.html',{'myset': dataset,'therapy':therapy})
            else:
                print("no data found")
    except Exception as e:
        print(f"Error->{e}")
    finally: 
        conn.close()

def heartbroken_therapy(req):
    conn=pymysql.connect(host='localhost',user='root',password='',database='psychbuddy',cursorclass=pymysql.cursors.DictCursor)
    dataset=[]
    therapy=None
    try:
        with conn.cursor() as cr:
            cr.execute("select * from sessions where sess_category=%s",('Heartbroken Therapy',))
            res=cr.fetchall()
            if res:
                for data in res:
                    therapy=data['sess_category']
                    dataset.append({
                        'name':data['sess_name'],
                        'category':data['sess_category'],
                        'path':data['sess_path'],
                        'desc':data['sess_desc']
                    })
                #print(dataset)
                return render(req,'sessionpage.html',{'myset': dataset,'therapy':therapy})
            else:
                print("no data found")
    except Exception as e:
        print(f"Error->{e}")
    finally: 
        conn.close()

def trauma_therapy(req):
    conn=pymysql.connect(host='localhost',user='root',password='',database='psychbuddy',cursorclass=pymysql.cursors.DictCursor)
    dataset=[]
    therapy=None
    try:
        with conn.cursor() as cr:
            cr.execute("select * from sessions where sess_category=%s",('Trauma Therapy',))
            res=cr.fetchall()
            if res:
                for data in res:
                    therapy=data['sess_category']
                    dataset.append({
                        'name':data['sess_name'],
                        'category':data['sess_category'],
                        'path':data['sess_path'],
                        'desc':data['sess_desc']
                    })
                print(dataset)
                return render(req,'sessionpage.html',{'myset': dataset,'therapy':therapy})
            else:
                print("no data found")
    except Exception as e:
        print(f"Error->{e}")
    finally: 
        conn.close()

def couple_therapy(req):
    conn=pymysql.connect(host='localhost',user='root',password='',database='psychbuddy',cursorclass=pymysql.cursors.DictCursor)
    dataset=[]
    therapy=None
    try:
        with conn.cursor() as cr:
            cr.execute("select * from sessions where sess_category=%s",('Couple Therapy',))
            res=cr.fetchall()
            if res:
                for data in res:
                    therapy=data['sess_category']
                    dataset.append({
                        'name':data['sess_name'],
                        'category':data['sess_category'],
                        'path':data['sess_path'],
                        'desc':data['sess_desc']
                    })
                print(dataset)
                return render(req,'sessionpage.html',{'myset': dataset,'therapy':therapy})
            else:
                print("no data found")
    except Exception as e:
        print(f"Error->{e}")
    finally: 
        conn.close()

def family_therapy(req):
    conn=pymysql.connect(host='localhost',user='root',password='',database='psychbuddy',cursorclass=pymysql.cursors.DictCursor)
    dataset=[]
    therapy=None
    try:
        with conn.cursor() as cr:
            cr.execute("select * from sessions where sess_category=%s",('Family Therapy',))
            res=cr.fetchall()
            if res:
                for data in res:
                    therapy=data['sess_category']
                    dataset.append({
                        'name':data['sess_name'],
                        'category':data['sess_category'],
                        'path':data['sess_path'],
                        'desc':data['sess_desc']
                    })
                print(dataset)
                return render(req,'sessionpage.html',{'myset': dataset,'therapy':therapy})
            else:
                print("no data found")
    except Exception as e:
        print(f"Error->{e}")
    finally: 
        conn.close()
#audio therapy---------------------------------------------------------------------------------------

def audio_therapy(req):
    conn = pymysql.connect(
        host='localhost',
        user='root',
        password='',
        database='psychbuddy',
        cursorclass=pymysql.cursors.DictCursor
    )
    mydata=[]
    try:
        with conn.cursor() as cr:
            cr.execute("select * from sound_therapy")
            rs=cr.fetchall()
            if rs: 
                for data in rs:
                    mydata.append({
                        'ther_name': data['therapy_name'],
                        'ther_img': data['therapy_img'],
                        'ther_path': data['therapy_path'],
                        'ther_desc': data['therapy_desc']
                    })
                #print(mydata)
    except Exception as e:
        print("Error -> ",str(e))
    finally: 
        conn.close()
    
    return render(req,'relaxation.html', {'dictionary': mydata})

#Feddback---------------------------------------------------------------------------------------------->

def feedback(req):
    conn=pymysql.connect(
        host="localhost",
        user="root",
        password="",
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
    mobileno=None
    try:
        with conn.cursor() as cr:
            cr.execute("select mobileno from users where login_status=%s",(1))
            res=cr.fetchone()
            if res:
                mobileno=res['mobileno']
                print(f"Moobile -> {mobileno}")
            if req.method=='POST':
                cr.execute("insert into feedback values(%s,%s,%s,%s)",(req.POST.get('tb1'),mobileno,req.POST.get('tb3'),None))
                res=conn.commit()
                return render(req,"feedback.html",{'res':"<script>alert('Feedback Successfull');</script>"})
    except Exception as e:
        print("Error ->",e)
    finally:
        conn.close()

    return render(req,"feedback.html")


#psychiatrist Information--------------------------------------------------------------------------------------------------->

def psychiatrist(req):
    conn=pymysql.connect(host='localhost',user='root',password='',database='psychbuddy',cursorclass=pymysql.cursors.DictCursor)
    data=[]
    try:
        with conn.cursor() as cr:
            cr.execute("select * from psychiatrist_dir")
            res=cr.fetchall()
            if res:
                for inf in res:
                    data.append({
                        'name':inf['psychname'],
                        'picture':inf['picture'],
                        'location':inf['location'],
                        'education':inf['education'],
                        'description':inf['description'],
                        'experience':inf['experience'],
                        'rating':inf['g-rating'],
                        'contact':inf['contact'],
                        'email':inf['email']
                    })
                #print(data)
    except Exception as e:
        print(f"Error -> {e}")
    
    finally:
        conn.close()

    return render(req,"psychiatrist.html",{'Data': data})

def psych_info(req):
    psychname = str(req.POST.get('psychname')).strip()
    print(f"Retrieved psychname: {psychname}")

    conn = pymysql.connect(
        host='localhost', user='root', password='', database='psychbuddy', cursorclass=pymysql.cursors.DictCursor
    )
    mydata = []
    useremail = None
    username = None

    try:
        global email
        global psychiatristname
        with conn.cursor() as cr2:
            cr2.execute("SELECT * FROM users WHERE login_status = %s", (1,))
            res2 = cr2.fetchone()
            if res2:
                useremail = res2['email']
                username = res2['fullname']
            #print(f"Logged-in user details: {username}, {useremail}")

        with conn.cursor() as cr:
            cr.execute("SELECT * FROM psychiatrist_dir WHERE psychname = %s", (psychname,))
            res = cr.fetchall()
            if res:
                for inf in res:
                    email = inf['email']
                    psychiatristname = inf['psychname']
                    mydata.append({
                        'name': inf['psychname'],
                        'picture': inf['picture'],
                        'location': inf['location'],
                        'education': inf['education'],
                        'description': inf['description'],
                        'experience': inf['experience'],
                        'rating': inf['g-rating'],
                        'email': inf['email'],
                        'contact': inf['contact']
                    })
                    break
                #print(f"Psychiatrist details: name: {psychiatristname}, email: {email}")
        if str(req.POST.get('appointmentBtn')) == "BookAppointment":
            yag = yagmail.SMTP('mypyschbuddy@gmail.com', 'pewcrroshhqeqalf')
            subject = 'Appointment Request via Psychbuddy Platform'
            body = f'''Dear Dr. {psychiatristname},

This is an appointment request submitted through the Psychbuddy platform by one of our users, {username}. They are seeking your expertise for a consultation regarding mental health concerns. Below are the details they have shared:

Name: {username}
Contact Information: {useremail}
Concern: I have been experiencing [e.g., stress, anxiety, mood swings, or difficulty coping with daily challenges] and am seeking professional guidance to understand and manage these concerns better. I would like to discuss my mental health challenges and explore potential treatment or coping strategies.
Additionally, the user has been maintaining a daily log of their mental state, which provides valuable insights into their emotional well-being and personality. This log is stored in a DOCX file and is attached to this email. We believe this information will assist you in better understanding the user’s ongoing emotional patterns and help guide the treatment plan.
Please let us know your availability or any next steps required to confirm the appointment. Should you need further details or clarification, feel free to reach out to the user directly or contact us at Psychbuddy.

Thank you for your attention and support.

Best regards,
The Psychbuddy Team
Contact: mypyschbuddy@gmail.com'''
            logpath = f"G:\\psychbuddy\\userlogs\\{username} Log.docx"
            yag.send(to=email, subject=subject, contents=body, attachments=logpath)
            return HttpResponse("<html><body><script>alert('Appointment Request Sent Successfully.'); window.location='psychiatrist';</script></body></html>")

    except Exception as e:
        print(f"Error -> {e}")
    finally:
        conn.close()

    return render(req, "psychInfo.html", {'myinfo': mydata})

#Reading Resources-------------------------------------------------------------------------------->

def books(req):
    conn=pymysql.connect(
        host="localhost",
        user="root",
        password="",
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
    data=[]
    try:
        with conn.cursor() as cr:
            cr.execute("select * from reading_resources",None)
            res=cr.fetchall()
            if res:
                for tr in res:
                    data.append({
                        'bname':tr['bookname'],
                        'bcover':tr['bookcover'],
                        'bauthor':tr['bookauthor'],
                        'bestfor':tr['bestfor'],
                        'bpyear':tr['bookpyear'],
                        'bpath':tr['bookpath'],
                        'bdesc':tr['bookdesc']
                    })
                #print(data)
    except Exception as e:
        print(f"Error -> {str(e)}") 
    finally:
        conn.close()
    
    return render(req,'books.html',{'mydata':data})

#Wellness Log----------------------------------------------------------------------------------->

def wellnessLog(req):
    fullname = None
    email = None
    conn = pymysql.connect(
        host='localhost',
        user='root',
        password=None,
        database='psychbuddy',
        cursorclass=pymysql.cursors.DictCursor
    )
    try:
        with conn.cursor() as cr:
            cr.execute("SELECT * FROM users WHERE login_status=%s", (1,))
            res = cr.fetchall()
            if res:
                for i in res:
                    fullname = i['fullname']
                    email = i['email']
            else:
                print("Error -> No user is active now!")
                return render(req, 'MentalWellnessLog.html', {'alert': "<script>alert('No active user found!');</script>"})
            if req.method == 'POST':
                cr.execute("insert into wellnesslog values(%s,%s,%s,%s)",(email, req.POST.get('tb1'),datetime.now().time(), date.today().strftime("%Y-%m-%d")))
                conn.commit()
                return render(req, 'MentalWellnessLog.html', {'name': fullname,'alert': "<script>alert('Data stored successfully!');</script>"
                })
            return render(req, 'MentalWellnessLog.html', {'name': fullname})
    except Exception as e:
        print("Error->",e)
        return render(req, 'MentalWellnessLog.html', {'alert': f"<script>alert('Error -> {e}');</script>"
        })
    finally:
        conn.close()

    return render(req, 'MentalWellnessLog.html')

# Wellness Logging Directory----------------------------------------------------------------------->

def wellnessLogDir(req):
    conn = pymysql.connect(
        host='localhost',
        user='root',
        password=None,
        database='psychbuddy',
        cursorclass=pymysql.cursors.DictCursor
    )
    rows = ['Name', 'Age', 'Address', 'Contact No', 'Email']
    udata = []
    mydata=[]
    no=0
    try:
        with conn.cursor() as cr:
            cr.execute("select * from users where login_status=%s", (1,))
            myres = cr.fetchone()
            if myres:
                udata = [
                    myres['fullname'],
                    myres['age'],
                    myres['address'],
                    myres['mobileno'],
                    myres['email']
                ]
            mixed=zip(rows,udata)
            cr.execute("select wellnesslog.* from wellnesslog join users on wellnesslog.email = users.email where users.login_status = %s",(1,))
            res2=cr.fetchall()
            if res2:
                for i in res2:
                    no=no+1
                    mydata.append({
                        'no':no,
                        'logmsg':i['log'],
                        'time':i['time'],
                        'date':i['date']
                    })
            return render(req, 'MentalWellnessLogDir.html', {'data': mixed,'data2':mydata})
    except Exception as e:
        print(f"Error -> {e}")
    finally:
        conn.close()

    return render(req, 'MentalWellnessLogDir.html')


#wellnesslogwordfilemaking---------------------------------------------------------------------->

def wellnessLogFileCreator(req):
    conn = pymysql.connect(
        host='localhost',
        user='root',
        password=None,
        database='psychbuddy',
        cursorclass=pymysql.cursors.DictCursor
    )
    mydata = []
    try:
        with conn.cursor() as cr:
            cr.execute("select wellnesslog.*, users.fullname, users.age, users.dob, users.mobileno, users.email as user_email from wellnesslog join users on wellnesslog.email = users.email where users.login_status = %s", (1,))
            res = cr.fetchall()
            if res:
                for i in res:
                    mydata.append({
                        'logmsg': i['log'],
                        'time': i['time'],
                        'date': i['date'],
                        'name': i['fullname'],
                        'age': i['age'],
                        'dob': i['dob'],
                        'mobile': i['mobileno'],
                        'email': i['user_email']
                    })
                mydoc = Document()
                mydoc.add_heading('Psychbuddy User Daily Log Record', level=0)
                username = res[0]['fullname']
                mydoc.add_paragraph(f"User: {username}\n")
                for log in mydata:
                    mydoc.add_paragraph(f"Log Entry\nDate: {log['date']} | Time: {log['time']}")
                    mydoc.add_paragraph(f"Log: {log['logmsg']}\n")
                file_path = f"G:\\psychbuddy\\userlogs\\{username} Log.docx"
                mydoc.save(file_path)
                #print(f"Document saved as {file_path}")

    except Exception as e:
        print(f"Error -> {e}")
    finally: 
        conn.close()
    return HttpResponse("<html><body><script>alert('Log file generated successfully!'); window.location='wellnesslogdir';</script></body></html>", content_type="text/html")

#yoga&meditationPage------------------------------------------------------------------------------>

def YandM(req):
    conn=pymysql.connect(
        host='localhost',
        user='root',
        password=None,
        database='psychbuddy',
        cursorclass=pymysql.cursors.DictCursor
    )
    data=[]
    data2=[]
    try:
        with conn.cursor() as cr:
            cr.execute("select * from yoga")
            res=cr.fetchall()
            if res:
                for rw in res:
                    data.append({
                        'r1':rw['yname'],
                        'r2':rw['steps'],
                        'r3':rw['benefits'],
                        'r4':rw['img']
                    })
                print(data)
            else:
                print("No data found inside table!")
            cr.execute("select * from meditation")
            res2=cr.fetchall()
            if res2:
                for rw2 in res2:
                    data2.append({
                        'name':rw2['mname'],
                        'desc':rw2['description']
                    })
                print(data2)
            else:
                print("No data found!")
    except Exception as e:
        print(f"Error -> {e}")
    
    return render(req,"YM.html",{'mydata': data,'mydata2': data2})

#emotion detecting module------------------------------------------------------------------------>
@csrf_exempt
def emotion_detection(request):
    if request.method == 'POST':
        emotion = detect_emo(duration=30)
        request.session['emotion']=emotion
        print(emotion)
        return JsonResponse({'emotion': emotion})
    
    return JsonResponse({'error': 'Invalid request'}, status=400)



#redirectors------------------------>

def homePage(req):
    try:
        emotion=detect_emo(duration=30)
        req.session['emotion']=emotion
        print(emotion)
    except Exception as e:
        print(f"Error -> {e}")
    return render(req,"homepage.html")

#personalized tips module--------------------------------------------------------------------->

def tips(req):
    conn=pymysql.connect(
        host='localhost',
        user='root',
        password=None,
        database='psychbuddy',
        cursorclass=pymysql.cursors.DictCursor
    )
    emotion=None
    suggestion=None
    emotions=req.session.get('emotion','neutral')
    try:
        cr=conn.cursor()
        cr.execute("select * from tips where emotion=%s",(emotions))
        res=cr.fetchone()
        if res:
            emotion=res['emotion']
            suggestion=res['suggestion']
            print(res['suggestion'])
    except Exception as e:
        print(f"Error->{e}")
    finally:
        conn.close()
    
    conn=pymysql.connect(
        host='localhost',
        user='root',
        password=None,
        database='psychbuddy',
        cursorclass=pymysql.cursors.DictCursor
    )
    admintips=None
    datetime=None
    try:
        cr=conn.cursor()
        cr.execute("select * from admintips")
        res=cr.fetchone()
        if res:
            admintips=res['tips']
            datetime=res['datetime']
    except Exception as e:
        print(f"Error->{e}")
    finally:
        conn.close()
    return render(req,"tips.html",{'emotion': emotion,'suggestion':suggestion,'tips': admintips, 'datetime': datetime}) 

def GIA(req):
    try:
        directory = r"G:/psychbuddy/GIA/"
        command = 'rasa run --enable-api --cors "*"'
        cmd = f'start cmd /k "cd /d {directory} && {command}"'
        subprocess.run(cmd, shell=True)
    except Exception as e:
        print(f"An error occurred: {e}")
    return render(req, "GIA.html")

def forum(req):
    return render(req,"forum.html")
def sessions(req):
    return render(req,"sessions.html")
def aboutus(req):
    return render(req,'aboutus.html')

def instagram(req):
    return HttpResponseRedirect("https://www.instagram.com/the_gt_official?utm_source=qr&igsh=MW10azA5aTBkaTYxbA==")

def whatsapp(req):
    return HttpResponseRedirect("https://wa.me/qr/TTCPPLO4MHO3B1")

def facebook(req):
    return HttpResponseRedirect("https://www.facebook.com/profile.php?id=100085688419302&mibextid=rS40aB7S9Ucbxw6v")

# Admin -------------------------------------------------------------------------------------------------------------------->

def adminLogin(req):
    if req.method == 'POST':
        adminname = str(req.POST.get('adminname'))
        pwd = str(req.POST.get('pass'))
        encoded_adminname = adminname.encode().hex()
        encoded_pwd = pwd.encode().hex()
        conn = pymysql.connect(
            host='localhost',
            user='root',
            password='',
            database='psychbuddy',
            cursorclass=pymysql.cursors.DictCursor
        )
        try:
            with conn.cursor() as cursor:
                sql_query = "SELECT * FROM admin_credentials WHERE adminname = %s AND password = %s"
                cursor.execute(sql_query, (encoded_adminname, encoded_pwd))
                rows = cursor.fetchall()
                if rows:
                    return HttpResponse("<html><body><script>window.location='admin';</script></body></html>")
                else:
                    return HttpResponse("<html><body><script>alert('Invalid Credentials Entered. Please Try Again!'); window.location='adminlogin';</script></body></html>")
        except Exception as e:
            messages.error(req, f"An error occurred: {e}")
        finally:
            conn.close()
    return render(req, "adminlogin.html")

# Admin Dashboard -------------------------------------------------------------------------------------------------------------------------------------------------------------->

def admindashboard(req):
    return render(req,"admindashboard.html")

def usersNav(req):
    conn=pymysql.connect(
        host='localhost',
        user='root',
        password=None,
        database='psychbuddy',
        cursorclass=pymysql.cursors.DictCursor
    )
    data=[]
    with conn.cursor() as cr:
        cr.execute("select * from users",(None))
        res=cr.fetchall()
        if res:
            for dt in res:
                data.append({
                    'username':dt['email'],
                    'fullname':dt['fullname'],
                    'contact':dt['mobileno'],
                    'email':dt['email']
                })
            #print(data)
        if req.method=='POST':
            email=req.POST.get('tb1')
            if email:
                cr.execute("delete from users where email=%s",(email,))
                conn.commit()
                return HttpResponse("<html><body><script>alert('User Account deleted successfully!');window.location='usersNav';</script></body></html>")
            else:
                return HttpResponse("<html><body><script>alert('Error Occured!');window.location='usersNav';</script></body></html>")
    return render(req,"usersNav.html",{'data':data})

def manageResources(req):
    return render(req,"manageResources.html")

def videoResourcesUpload(rq):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
    if rq.method == "POST":
        file = rq.FILES.get('sessionfile')
        sessname = rq.POST.get('sessionname')
        sessdesc = rq.POST.get('sessiondesc')
        sesscategory = rq.POST.get('sessioncategory')
        path = None
        filename = file.name if file else None

        try:
            if file:
                if sesscategory == "Individual Therapy":
                    path = f"G:/psychbuddy/static/sessions/individual therapy/{filename}"
                elif sesscategory == "Group Therapy":
                    path = f"G:/psychbuddy/static/sessions/group therapy/{filename}"
                elif sesscategory == "Couple Therapy":
                    path = f"G:/psychbuddy/static/sessions/couple therapy/{filename}"
                elif sesscategory == "Trauma Therapy":
                    path = f"G:/psychbuddy/static/sessions/trauma therapy/{filename}"
                elif sesscategory == "HeartBroken Therapy":
                    path = f"G:/psychbuddy/static/sessions/heartbroken therapy/{filename}"
                else:
                    path = f"G:/psychbuddy/static/sessions/family therapy/{filename}"

                if not os.path.exists(os.path.dirname(path)):
                    os.makedirs(os.path.dirname(path))

                dbsesspath = path[21:]  
                with conn.cursor() as cr:
                    query = "insert into sessions values(%s, %s, %s, %s)"
                    cr.execute(query, (sessname, sesscategory, dbsesspath, sessdesc))
                    shutil.copy(file.temporary_file_path(), path)
                    conn.commit()
                    return HttpResponse(
                        "<html><body><script>alert('Task Completed Successfully.');"
                        "window.location='videoRes';</script></body></html>"
                    )
            else:
                return HttpResponse(
                    "<html><body><script>alert('No file uploaded!');"
                    "window.location='videoRes';</script></body></html>"
                )
        except Exception as e:
            print("Error->", e)
            return HttpResponse(
                f"<html><body><script>alert('Error: {str(e)}');"
                "window.location='videoRes';</script></body></html>"
            )
        finally:
            conn.close()
    return render(rq, "videoResAdmin.html")

def audioResourcesUpload(req):
    conn = pymysql.connect(
    host="localhost",
    user="root",
    password=None,
    database="psychbuddy",
    cursorclass=pymysql.cursors.DictCursor
    )
    if req.method == "POST":
        file = req.FILES.get('audiofilepath')
        file2=req.FILES.get('audioimgpath')
        audioname = req.POST.get('audioname')
        audiodesc = req.POST.get('audiodesc')
        path = None #audiimgopath
        path2=None #audiopath
        filename = file.name
        filename2=file2.name

        try:
            if file and file2:
                path=f"G:/psychbuddy/static/img/sound_img/{filename2}"
                path2=f"G:/psychbuddy/static/audio/{filename}"

                dbaudioimgpath = path[21:]
                dbaudiopath=path2[21:]  
                with conn.cursor() as cr:
                    query = "insert into sound_therapy values(%s, %s, %s, %s)"
                    cr.execute(query, (audioname, dbaudioimgpath, dbaudiopath, audiodesc))
                    with open(path, 'wb') as f_img:
                        shutil.copyfileobj(file2, f_img) #tempory_file_path() vannot use for small files

                    with open(path2, 'wb') as f_audio:
                        shutil.copyfileobj(file, f_audio)

                    conn.commit()
                    return HttpResponse(
                        "<html><body><script>alert('Tasking Complete.');"
                        "window.location='videoRes';</script></body></html>"
                    )
            else:
                return HttpResponse(
                    "<html><body><script>alert('No file uploaded, Something gone wrong!');"
                    "window.location='videoRes';</script></body></html>"
                )
        except Exception as e:
            print("Error->", e)
            return HttpResponse(
                f"<html><body><script>alert('Error: {str(e)}');"
                "window.location='videoRes';</script></body></html>"
            )
        finally:
            conn.close()
    return render(req, "audioResAdmin.html")

def bookResourcesUpload(req):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
    if req.method == "POST":
        book_cover = req.FILES.get('bcover')  
        book_pdf = req.FILES.get('bfile')    
        book_name = req.POST.get('bname') 
        book_author = req.POST.get('bauthor')  
        best_for = req.POST.get('bestfor')
        publish_year = req.POST.get('byear')  
        book_desc = req.POST.get('audiodesc') 

        cover_path = None
        pdf_path = None

        try:
            if book_cover and book_pdf:
                cover_filename = book_cover.name
                pdf_filename = book_pdf.name
                cover_path = f"G:/psychbuddy/static/img/books_poster/{cover_filename}"
                pdf_path = f"G:/psychbuddy/static/books/{pdf_filename}"
                db_cover_path = cover_path[21:]
                db_pdf_path = pdf_path[21:]

                with conn.cursor() as cr:
                    query = "INSERT INTO reading_resources VALUES (%s, %s, %s, %s, %s, %s, %s)"
                    cr.execute(query, (
                        book_name, db_cover_path, book_author, best_for, publish_year, db_pdf_path, book_desc
                    ))
                    with open(cover_path, 'wb') as f_cover:
                        shutil.copyfileobj(book_cover, f_cover)
                    with open(pdf_path, 'wb') as f_pdf:
                        shutil.copyfileobj(book_pdf, f_pdf)
                    conn.commit()
                    return HttpResponse(
                        "<html><body><script>alert('Book Resource Uploaded Successfully.');"
                        "window.location='bookRes';</script></body></html>"
                    )
            else:
                return HttpResponse(
                    "<html><body><script>alert('Please upload both book cover and PDF file.');"
                    "window.location='bookRes';</script></body></html>"
                )
        except Exception as e:
            print("Error->", e)
            return HttpResponse(
                f"<html><body><script>alert('Error: {str(e)}');"
                "window.location='bookRes';</script></body></html>"
            )
        finally:
            conn.close()

    return render(req, "readingResAdmin.html")


def yResourcesUpload(req):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
    if req.method == "POST":
        if req.POST.get('yname'):
            yoga_name = req.POST.get('yname')
            yoga_steps = req.POST.get('steps')
            yoga_benefits = req.POST.get('ydesc')
            yoga_image = req.FILES.get('yimgpath')

            if yoga_image:
                yoga_image_path = f"G:/psychbuddy/static/img/YM/{yoga_image.name}"
                db_yoga_image_path = yoga_image_path[21:]

                try:
                    with conn.cursor() as cr:
                        query = "insert into yoga values(%s, %s, %s, %s)"
                        cr.execute(query, (yoga_name, yoga_steps, yoga_benefits,db_yoga_image_path ))

                        with open(yoga_image_path, 'wb') as f_yoga_img:
                            shutil.copyfileobj(yoga_image, f_yoga_img)

                        conn.commit()
                        return HttpResponse(
                            "<html><body><script>alert('Yoga Resource Uploaded Successfully.');"
                            "window.location='yRes';</script></body></html>"
                        )
                except Exception as e:
                    print("Error->", e)
                    return HttpResponse(
                        f"<html><body><script>alert('Error: {str(e)}');"
                        "window.location='ymRes';</script></body></html>"
                    )
                finally:
                    conn.close()
    return render(req, "yResAdmin.html")

def mResourcesUpload(req):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
    if req.method == "POST":
        meditation_name = req.POST.get('Mname')
        meditation_desc = req.POST.get('Mdesc')
        try:
            with conn.cursor() as cr:
                query = "insert into meditation values(%s, %s)"
                cr.execute(query, (meditation_name, meditation_desc))
                conn.commit()
                return HttpResponse("<html><body><script>alert('Meditation Resource Uploaded Successfully.'); window.location='mRes';</script></body></html>")
        except Exception as e:
            print("Error->", e)
        finally:
            conn.close()
    return render(req,"mResAdmin.html")

def deleteYoga(request):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )

    if request.method == "POST":
        yoga_name = request.POST.get('yogaName', '').strip()

        if yoga_name:
            try:
                with conn.cursor() as cr:
                    cr.execute("delete from yoga where yname=%s", (yoga_name,))
                    conn.commit()
                    return HttpResponse("<html><body><script>alert('Record Deleted!'); window.location='delYRes';</script></body></html>")
            except Exception as e:
                print(f"Error: {e}")
            finally:
                conn.close()
    return render(request, "delYRes.html")

def deleteMeditation(request):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )

    if request.method == "POST":
        meditation_name = request.POST.get('meditationName')

        if meditation_name:
            try:
                with conn.cursor() as cr:
                    cr.execute("delete from meditation where mname=%s", (meditation_name,))
                    conn.commit()
                    return HttpResponse("<html><body><script>alert('Record Deleted!'); window.location='delMRes';</script></body></html>")
            except Exception as e:
                print(f"Error: {e}")
            finally:
                conn.close()

    return render(request, "delMRes.html")



def deleteBookResource(request):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )

    if request.method == "POST":
        book_name = request.POST.get('bname').strip()

        if book_name:
            try:
                with conn.cursor() as cr:
                    check_query = "select * from reading_resources where bookname = %s"
                    cr.execute(check_query, (book_name,))
                    result = cr.fetchone()

                    if result:
                        delete_query = "delete from reading_resources where bookname = %s"
                        cr.execute(delete_query, (book_name,))
                        conn.commit()
                        return HttpResponse(
                            "<html><body><script>alert('Book Resource Deleted Successfully.');"
                            "window.location='/adminPanel';</script></body></html>"
                        )
                    else:
                        return HttpResponse(
                            "<html><body><script>alert('Book Resource Not Found.');"
                            "window.location='/adminPanel';</script></body></html>"
                        )

            except Exception as e:
                print("Error->", e)
                return HttpResponse(
                    f"<html><body><script>alert('Error: {str(e)}');"
                    "window.location='/adminPanel';</script></body></html>"
                )
            finally:
                conn.close()

    return render(request, "delBookRes.html")


def deleteAudioResource(request):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )

    if request.method == "POST":
        audio_name = request.POST.get('audioname').strip()

        if audio_name:
            try:
                with conn.cursor() as cr:
                    check_query = "SELECT * FROM sound_therapy WHERE therapy_name = %s"
                    cr.execute(check_query, (audio_name,))
                    result = cr.fetchone()

                    if result:
                        delete_query = "DELETE FROM sound_therapy WHERE therapy_name = %s"
                        cr.execute(delete_query, (audio_name,))
                        conn.commit()
                        messages.success(request, "Audio Resource Deleted Successfully.")
                        return redirect('resources')
                    else:
                        messages.error(request, "Audio Resource Not Found.")
                        return redirect('resources')

            except Exception as e:
                messages.error(request, f"Error: {str(e)}")
                return redirect('resources')

            finally:
                conn.close()

    return render(request, "delAudioRes.html")

def deleteVideoResource(request):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password=None,
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )

    if request.method == "POST":
        video_name = request.POST.get('videoname').strip()

        if video_name:
            try:
                with conn.cursor() as cr:
                    check_query = "select * from sessions where sess_name = %s"
                    cr.execute(check_query, (video_name,))
                    result = cr.fetchone()

                    if result:
                        delete_query = "delete from sessions where sess_name = %s"
                        cr.execute(delete_query, (video_name,))
                        conn.commit()
                        messages.success(request, "Video Resource Deleted Successfully.")
                        return redirect('resources')
                    else:
                        messages.error(request, "Video Resource Not Found.")
                        return redirect('resources')

            except Exception as e:
                messages.error(request, f"Error: {str(e)}")
                return redirect('resources')

            finally:
                conn.close()

    return render(request, "delVideoRes.html")

def chatAdminForum(request):
    conn = pymysql.connect(
        host="localhost",
        user="root",
        password="",
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
    usermsg = []
    adminmsg=[]

    try:
        with conn.cursor() as cr:
            sql = """SELECT users.fullname, community_forum.date, community_forum.time, community_forum.msg FROM users JOIN community_forum ON community_forum.email = users.email WHERE community_forum.administrator IS Null ORDER BY community_forum.date ASC, community_forum.time ASC;
"""
            cr.execute(sql)
            res = cr.fetchall()
            if res:
                for data in res:
                    usermsg.append({
                        'msg': data['msg'],
                        'cdate': data['date'],
                        'ctime': data['time'],
                        'name': data['fullname']
                    })
                print(usermsg)
    except Exception as e:
        print(f"Error: {e}")
    finally:
        conn.close()
    
    try:
        conn = pymysql.connect(
        host="localhost",
        user="root",
        password="",
        database="psychbuddy",
        cursorclass=pymysql.cursors.DictCursor
    )
        with conn.cursor() as cr:
            sql = "select * from community_forum where administrator = %s order by date asc, time asc"
            cr.execute(sql, ('GT'))
            res = cr.fetchall()
            if res:
                for data in res:
                    adminmsg.append({
                        'msg': data['msg'],
                        'cdate': data['date'],
                        'ctime': data['time'],
                    })
    except Exception as e:
        print(f"Error: {e}")
    finally:
        conn.close()

    if request.method == 'POST':
        msg = request.POST.get('msg')
        conn = pymysql.connect(
            host='localhost',
            user='root',
            password='',
            database='psychbuddy',
            cursorclass=pymysql.cursors.DictCursor
        )
        try:
            with conn.cursor() as cr:
                sql2 = """INSERT INTO community_forum (msg, email, date, time, administrator)
VALUES (%s, %s, %s, %s, %s)"""
                cr.execute(sql2, (msg,'ganeshtelore4@gmail.com', date.today(), datetime.now().time(),'GT'))
                conn.commit()
                return JsonResponse({'success': True, 'msg': msg})
        except Exception as e:
            print(e)
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
        finally:
            conn.close()
    return render(request, "adminforum.html", {'info': usermsg,'info2':adminmsg})

def feedbackinfo(req):
    conn = pymysql.connect(
            host='localhost',
            user='root',
            password='',
            database='psychbuddy',
            cursorclass=pymysql.cursors.DictCursor
        )
    feedbacks=[]
    try:
        with conn.cursor() as cr:
            cr.execute("select * from feedback")
            rs=cr.fetchall()
            if rs:
                for i in rs:
                    feedbacks.append({
                        'name': i['user'],
                        'contact': i['contact'],
                        'msg': i['msg'],
                        'datetime': i['datetime']
                    })
    except Exception as e:
        print(f"Error-> {e}")
    finally: 
        conn.close()
    return render(req,"feedbackinfo.html",{'data': feedbacks})

def fixappointmentmail(req):
    if req.method=='POST':
        email=req.POST.get('tb1')
        #desc=req.POST.get('tb2')
        appointtime=req.POST.get('tb3')
        appointdate=req.POST.get('tb4')
        try:
            yag = yagmail.SMTP('mypyschbuddy@gmail.com', 'pewcrroshhqeqalf')
            subject = 'Appointment Confirmation - Psychbuddy'
            body = f'''Dear Psychbuddy User,

We are pleased to inform you that your appointment request with the psychiatrist through Psychbuddy has been processed. Below are the details of your appointment:

Appointment Date: {appointdate[::-1]}
Appointment Time: {appointtime}

Thank you for choosing Psychbuddy for your mental health journey. Should you need further assistance, feel free to reach out to us.

Warm regards,
Psychbuddy Team'''
            yag.send(to=email, subject=subject, contents=body)
            return HttpResponse("<script>alert('Mail Sent.');window.location='appointmentmail';</script>")
        except Exception as e:
            print(f"Error -> {e}")
        
    return render(req,"appointmentmail.html")

def admintips(req):
    conn = pymysql.connect(
            host='localhost',
            user='root',
            password='',
            database='psychbuddy',
            cursorclass=pymysql.cursors.DictCursor
        )
    if req.method=="POST":
        tips=req.POST.get('tip')
        try:
            with conn.cursor() as cr:
                cr.execute("insert into admintips values(%s,%s)",(tips,None))
                conn.commit()
                return HttpResponse("<script>alert('Tips Appended.'); window.location='admintips';</script>")
        except Exception as e:
            print(f"Error -> {e}")
        finally:
            conn.close()
    return render(req,"admintips.html")

def adminlogout(req):
    return render(req,"homepage.html")